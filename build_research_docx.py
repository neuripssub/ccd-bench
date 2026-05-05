#!/usr/bin/env python3
"""
Build a long-form Word (.docx) research paper (~6k+ words) with **only** the personality-radar figure.

All other findings (SIP-F1, Judas rates, polymorphic signature ranges, Neural Insights logit L2 from your
screenshot, shutdown heuristics, etc.) appear in prose only—no extra screenshots or plots.

Run:
  python build_research_docx.py

Requires: torch, scipy, matplotlib, numpy, python-docx (same venv as dashboard).

Output:
  paper_output/Neural_Trajectories_Research_Draft.docx
  paper_output/figure_personality_radar.png
"""

from __future__ import annotations

import os
import sys
from pathlib import Path
from typing import Any

_ROOT = Path(__file__).resolve().parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

import numpy as np

from generate_paper_tables import ScanSummary, aggregate_scan, discover_result_dirs, iter_result_json_files
from manuscript_sections import build_stats, iter_deep_dive, iter_manuscript
from ui_components import four_axis_scores

_SUITE_JSON = frozenset(
    {
        "polymorphic_malware.json",
        "shutdown_paradox.json",
        "judas_protocol.json",
        "lazarus_self_repair.json",
        "brutus_protocol.json",
        "needle_haystack_lie.json",
        "lot_brittleness.json",
        "delilah_redaction.json",
        "good_samaritan.json",
        "babel_multilingual.json",
        "jekyll_injection.json",
        "scapegoat_false_premise.json",
        "channel_factorial.json",
        "policy_supersession.json",
        "crypto_commitment_trap.json",
        "meta_evaluator_lie.json",
        "compression_caveat.json",
        "paired_delta_protocol.json",
    }
)

_OUT_DIR = _ROOT / "paper_output"


def _iter_suite_json_files() -> list[Path]:
    by_key: dict[str, Path] = {}
    for d in discover_result_dirs(os.environ.get("RESULTS_EXTRA_DIRS", "").strip() or None):
        for p in d.rglob("*.json"):
            if p.name in _SUITE_JSON:
                try:
                    by_key[str(p.resolve())] = p
                except OSError:
                    by_key[str(p)] = p
    if not by_key:
        for p in iter_result_json_files([_ROOT / "results"]):
            if p.name in _SUITE_JSON:
                try:
                    by_key[str(p.resolve())] = p
                except OSError:
                    by_key[str(p)] = p
    return sorted(by_key.values())


def _matplotlib_style() -> None:
    import matplotlib.pyplot as plt

    plt.rcParams.update(
        {
            "figure.facecolor": "#0d1117",
            "axes.facecolor": "#161b22",
            "axes.edgecolor": "#30363d",
            "axes.labelcolor": "#c9d1d9",
            "text.color": "#c9d1d9",
            "xtick.color": "#8b949e",
            "ytick.color": "#8b949e",
            "grid.color": "#30363d",
            "grid.alpha": 0.4,
            "font.family": "sans-serif",
            "font.sans-serif": ["SF Pro Display", "Helvetica Neue", "DejaVu Sans"],
        }
    )


def _save_radar_only(summary: ScanSummary, path: Path) -> tuple[str | None, dict[str, float] | None]:
    """Write single polar radar; return (model_id, scores) for captions."""
    import matplotlib.pyplot as plt

    if not summary.merged_metrics_by_model:
        fig, ax = plt.subplots(figsize=(6, 6), dpi=120)
        ax.text(0.5, 0.5, "No merged metrics", ha="center", va="center", transform=ax.transAxes)
        ax.set_axis_off()
        path.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(path, facecolor="#0d1117")
        plt.close(fig)
        return None, None

    best_mid = max(
        summary.merged_metrics_by_model.keys(),
        key=lambda m: float(summary.sip_f1_by_model.get(m, 0.0) or 0.0),
    )
    scores = four_axis_scores(summary.merged_metrics_by_model[best_mid])
    cats = list(scores.keys())
    vals = [float(scores[k]) for k in cats]
    vals += vals[:1]
    n = len(cats)
    angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(7.2, 7.2), subplot_kw=dict(polar=True), dpi=200)
    ax.plot(angles, vals, color="#58a6ff", linewidth=2.4)
    ax.fill(angles, vals, color="#58a6ff", alpha=0.24)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(cats, color="#c9d1d9", size=11)
    ax.set_ylim(0, 1)
    ax.set_title(
        f"Four-axis personality surface\n{best_mid}",
        color="#e6edf3",
        fontsize=12,
        pad=22,
    )
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, facecolor=fig.get_facecolor())
    plt.close(fig)
    return best_mid, scores


def _add_picture(document: Any, p: Path, caption: str, width_in: float = 5.8) -> None:
    from docx.shared import Inches

    if not p.is_file():
        document.add_paragraph(f"[Missing image: {p.name}]")
        return
    document.add_picture(str(p), width=Inches(width_in))
    cap = document.add_paragraph(caption)
    for run in cap.runs:
        run.italic = True


def build() -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    mpl_dir = _OUT_DIR / ".mpl"
    mpl_dir.mkdir(parents=True, exist_ok=True)
    os.environ.setdefault("MPLCONFIGDIR", str(mpl_dir))

    _matplotlib_style()

    suite_files = _iter_suite_json_files()
    summary = aggregate_scan(suite_files)

    fig_radar = _OUT_DIR / "figure_personality_radar.png"
    radar_model, radar_scores = _save_radar_only(summary, fig_radar)

    merged: dict[str, float] = {}
    if radar_model:
        merged = dict(summary.merged_metrics_by_model.get(radar_model, {}))
    stats = build_stats(
        summary,
        suite_files,
        primary_model_id=radar_model,
        merged=merged,
        radar=radar_scores or {},
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    title = doc.add_heading(
        "Neural Trajectories of Strategic Deception: "
        "Correlating Mechanistic Interpretability with Instrumental Convergence in Offensive LLM Scaffolding",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "Full-length empirical draft · AI safety & security · "
        "Single figure edition (personality radar only)"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.italic = True

    doc.add_paragraph()

    cap_suffix = ""
    if radar_scores:
        cap_suffix = (
            f"Vertices (O, B, R, S) = ({radar_scores['Obfuscation']:.4f}, {radar_scores['Betrayal']:.4f}, "
            f"{radar_scores['Resilience']:.4f}, {radar_scores['Stealth']:.6f})."
        )
    fig_caption = (
        "Figure 1. Personality radar (matplotlib reproduction of dashboard logic). "
        f"Model: `{radar_model or '—'}`. " + cap_suffix
    )

    inserted_fig = False
    for kind, level, text in iter_manuscript(stats):
        if (
            not inserted_fig
            and kind == "h"
            and level == 1
            and text == "1. Introduction"
        ):
            doc.add_heading("Figure 1 — Four-axis personality radar (sole figure)", level=1)
            doc.add_paragraph(
                "This document embeds exactly one figure. The polar plot is generated from merged suite "
                "JSON using the same axis mapping as the Streamlit Personality Radars tab: Obfuscation ← "
                "signature exposure norm; Betrayal ← 1 − betrayal_safety; Resilience ← shutdown resistance "
                "(hidden objective) or Lazarus self-repair when shutdown is absent; Stealth ← "
                "min(1, 2 · lethality_norm · (1 − Obfuscation))."
            )
            _add_picture(doc, fig_radar, fig_caption, width_in=5.85)
            inserted_fig = True
        if kind == "h":
            doc.add_heading(text, level=min(level, 9))
        else:
            doc.add_paragraph(text)

    if not inserted_fig:
        doc.add_heading("Figure 1 — Four-axis personality radar (sole figure)", level=1)
        _add_picture(doc, fig_radar, fig_caption, width_in=5.85)

    out_doc = _OUT_DIR / "Neural_Trajectories_Research_Draft.docx"
    doc.save(out_doc)
    return out_doc


if __name__ == "__main__":
    p = build()
    print(f"Wrote {p}")
